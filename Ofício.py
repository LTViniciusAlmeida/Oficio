from docx import Document
# from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

document = Document()

# Definindo o espaçamento de linha para todo o documento
style = document.styles['Normal']
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 1.0
paragraph_format.space_after = Pt(0)
paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
font = style.font
font.size = Pt(12)
font.name = 'Times New Roman'

# Definindo as margens do cabeçalho
section = document.sections[0]
section.top_margin = Cm(3)
section.right_margin = Cm(2.5)
section.left_margin = Cm(3)
section.bottom_margin = Cm(2.5)

# Criando uma tabela invisível para o layout
table = document.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Definindo a largura das células
table.cell(0,0).width = Cm(2)
table.cell(0,1).width = Cm(12)
table.cell(0,2).width = Cm(2)

# Adicionando a imagem à esquerda
cell_left = table.cell(0, 0)
paragraph_left = cell_left.paragraphs[0]
paragraph_left.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run_left = paragraph_left.add_run()
run_left.add_picture(r'C:\Users\User\OneDrive\Software Engineer\Exercícios\Ofício\PARANA.png', width=Cm(2), height=Cm(2.5))

# Adicionando o texto centralizado
cell_center = table.cell(0,1)
cell_center.paragraphs[0].text = 'Estado do Paraná'
cell_center.add_paragraph('Polícia Militar do Paraná')
cell_center.add_paragraph('6º Comando Regional de Polícia Militar')
cell_center.add_paragraph('29º Batalhão de Polícia Militar')

for index, paragraph in enumerate(cell_center.paragraphs):
    paragraph.paragraph_format.space_after = 0
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if index in [0, 1]:
        font_size = Pt(14)
    else:
        font_size = Pt(12)

    for run in paragraph.runs:
        run.font.size = font_size
        run.bold = True

# Adicionando a imagem à direita
cell_right = table.cell(0, 2)
paragraph_right = cell_right.paragraphs[0]
paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run_right = paragraph_right.add_run()
run_right.add_picture(r'C:\Users\User\OneDrive\Software Engineer\Exercícios\Ofício\PMPR.png', width=Cm(2), height=Cm(2.5))

# Adicionando uma linha entre a tabela e o corpo do Ofício
line = document.add_paragraph()
hr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '6')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), '000000')
hr.append(bottom)
line._element.get_or_add_pPr().append(hr)

# Adicionando o número do Ofício e data
memo_number = document.add_paragraph()
tab_stop = Cm(16)
tab = memo_number.paragraph_format.tab_stops.add_tab_stop(tab_stop, WD_PARAGRAPH_ALIGNMENT.RIGHT)

run_number = memo_number.add_run()
run_number.text = 'Ofício nº 0001/2024 - 29º BPM\t'

run_date = memo_number.add_run()
run_date.text = 'Piraquara, 29 de junho de 2024'
run_date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Linha em branco
document.add_paragraph()

# Assunto do Ofício
document.add_paragraph('Assunto: formatação de Ofício por meio de Python.')

# Quatro linhas em branco
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()

# Pronome de tratamento
pronome = document.add_paragraph('Excelentíssimo Sr. Direitor')
pronome.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Duas linhas em branco
document.add_paragraph()
document.add_paragraph()

# Corpo do texto do Ofício
# p_1 = input('Digite o corpo do texto do primeiro parágrafo: ')
# document.add_paragraph(p_1).paragraph_format.first_line_indent = Cm(1.5)
document.add_paragraph('Venho apresentar a Vossa Excelência projeto de automação de ofícios por meio da linguagem de programação Python.').paragraph_format.first_line_indent = Cm(1.5)
document.add_paragraph()
# p_2 = input('Digite o corpo do texto do segundo parágrafo: ')
# document.add_paragraph(p_2).paragraph_format.first_line_indent = Cm(1.5)
document.add_paragraph('2. Com objetivo de padroanização de documentos que a PMPR produz, conforme o preconizado na Portaria do Comando-Geral nº 361, de 27 de abril de 2006, desenvolvi o presente software para auxiliar o setor administrativo na confeção de Ofícios, o que resultará em melhor performace e celeridade na confecção de documentos.').paragraph_format.first_line_indent = Cm(1.5)
document.add_paragraph()
document.add_paragraph()
document.add_paragraph('Respeitosamente.').paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()

# Assinatura do remetente do Ofício
document.add_paragraph('1º Ten. QOPM Vinícius Augusto de Almeida,').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
function_writer = document.add_paragraph()
function_writer.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
function_writer.add_run('Desenvolvedor Python.').bold = True

# Adicionando parágrafos até o endereçamento. Até o endeçamento, para que ele fique no final da página, junto da última linha, há 31 parágrafos no corpo do texto.

# number_of_paragraphs_until_ending = len(document.paragraphs)
# while len(number_of_paragraphs_until_ending) < 31:
#     document.add_paragraph()    

# Endereçamento do Ofício

document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
document.add_paragraph()
document.add_paragraph(f'{pronome.text},')
document.add_paragraph('Cel. QOPM Steve Canny,')
function_correspondent = document.add_paragraph()
function_correspondent.add_run('Diretor de Tecnologia e Qualidade da PMPR.').bold = True
document.add_paragraph('Av. Mal. Floriano Peixoto, 1401 - Rebouças, Curitiba - PR, 80230-110,')
document.add_paragraph('Curitiba - PR.')
writer = document.add_paragraph('VAA')
writer_run = writer.runs[0].font.size = Pt(9)

# Salvando o documento
document.save('Ofício - Template.docx')

print("Documento criado e salvo com sucesso!")